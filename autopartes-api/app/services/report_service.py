import io
from datetime import datetime, timezone
from decimal import Decimal
from sqlalchemy.orm import Session
from sqlalchemy import func
from app.models.order import Order, OrderItem, OrderStatus
from app.models.part import Part
from app.models.user import User, UserRole


# ═══════════════════════════════════════════════════════════════
# REPORTE 1 — Resumen general de ventas
# ═══════════════════════════════════════════════════════════════

def get_summary(db: Session) -> dict:
    """Totales de ventas, top 5 partes y pedidos por estado."""
    result = (
        db.query(
            func.coalesce(func.sum(Order.total_amount), 0).label("total_sales"),
            func.count(Order.id).label("total_orders"),
        )
        .filter(Order.status != OrderStatus.CANCELLED)
        .first()
    )

    top_rows = (
        db.query(
            Part.id.label("part_id"),
            Part.name,
            Part.sku,
            func.sum(OrderItem.quantity).label("total_quantity_sold"),
            func.sum(OrderItem.subtotal).label("total_revenue"),
        )
        .join(OrderItem, Part.id == OrderItem.part_id)
        .join(Order, Order.id == OrderItem.order_id)
        .filter(Order.status != OrderStatus.CANCELLED)
        .group_by(Part.id, Part.name, Part.sku)
        .order_by(func.sum(OrderItem.quantity).desc())
        .limit(5)
        .all()
    )

    top_parts = [
        {
            "part_id": str(row.part_id),
            "name": row.name,
            "sku": row.sku,
            "total_quantity_sold": row.total_quantity_sold,
            "total_revenue": Decimal(str(row.total_revenue)),
        }
        for row in top_rows
    ]

    status_rows = (
        db.query(Order.status, func.count(Order.id).label("count"))
        .group_by(Order.status)
        .all()
    )

    orders_by_status = {row.status.value: row.count for row in status_rows}
    for status in OrderStatus:
        orders_by_status.setdefault(status.value, 0)

    return {
        "total_sales": Decimal(str(result.total_sales)),
        "total_orders": result.total_orders,
        "top_parts": top_parts,
        "orders_by_status": orders_by_status,
    }


# ═══════════════════════════════════════════════════════════════
# REPORTE 2 — Pedidos detallado
# ═══════════════════════════════════════════════════════════════

def get_orders_report(db: Session) -> dict:
    """Lista detallada de pedidos con conteo de items."""
    orders = (
        db.query(Order)
        .order_by(Order.created_at.desc())
        .limit(200)
        .all()
    )
    return {
        "orders": [
            {
                "id": str(o.id),
                "user_id": str(o.user_id),
                "status": o.status.value,
                "total_amount": Decimal(str(o.total_amount)),
                "items_count": len(o.items),
                "created_at": o.created_at.isoformat() if o.created_at else None,
            }
            for o in orders
        ],
        "total": len(orders),
    }


# ═══════════════════════════════════════════════════════════════
# REPORTE 3 — Clientes por volumen de compras
# ═══════════════════════════════════════════════════════════════

def get_clients_report(db: Session) -> dict:
    """Clientes externos ordenados por total gastado."""
    rows = (
        db.query(
            User.id.label("user_id"),
            User.full_name,
            User.email,
            func.count(Order.id).label("total_orders"),
            func.coalesce(func.sum(Order.total_amount), 0).label("total_spent"),
        )
        .outerjoin(Order, (Order.user_id == User.id) & (Order.status != OrderStatus.CANCELLED))
        .filter(User.role == UserRole.CUSTOMER)
        .group_by(User.id, User.full_name, User.email)
        .order_by(func.coalesce(func.sum(Order.total_amount), 0).desc())
        .all()
    )
    return {
        "clients": [
            {
                "user_id": str(r.user_id),
                "full_name": r.full_name,
                "email": r.email,
                "total_orders": r.total_orders,
                "total_spent": Decimal(str(r.total_spent)),
            }
            for r in rows
        ],
        "total_clients": len(rows),
    }


# ═══════════════════════════════════════════════════════════════
# REPORTE 4 — Estado actual del inventario
# ═══════════════════════════════════════════════════════════════

def get_inventory_report(db: Session) -> dict:
    """Stock actual de todas las autopartes activas. Marca low_stock <= 5."""
    parts = (
        db.query(Part)
        .filter(Part.is_active == True)
        .order_by(Part.stock_quantity.asc())
        .all()
    )
    return {
        "parts": [
            {
                "id": str(p.id),
                "sku": p.sku,
                "name": p.name,
                "category": p.category,
                "stock_quantity": p.stock_quantity,
                "price": Decimal(str(p.price)),
                "low_stock": p.stock_quantity <= 5,
            }
            for p in parts
        ],
        "total_parts": len(parts),
        "low_stock_count": sum(1 for p in parts if p.stock_quantity <= 5),
    }


# ═══════════════════════════════════════════════════════════════
# GENERACIÓN DE ARCHIVOS — helpers internos
# ═══════════════════════════════════════════════════════════════

def _now_str() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _pdf_header_style():
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle
    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#ecf0f1")]),
    ])


# ── PDF: Resumen de ventas ────────────────────────────────────

def generate_pdf(data: dict) -> bytes:
    """PDF del resumen general de ventas."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Reporte de Ventas — Autopartes API", styles["Title"]))
    elements.append(Paragraph(f"Generado: {_now_str()}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    elements.append(Paragraph("Resumen General", styles["Heading2"]))
    summary_data = [
        ["Métrica", "Valor"],
        ["Total de Ventas (excl. cancelados)", f"${float(data['total_sales']):,.2f}"],
        ["Total de Pedidos (excl. cancelados)", str(data["total_orders"])],
    ]
    t = Table(summary_data, colWidths=[3.5 * inch, 3 * inch])
    t.setStyle(_pdf_header_style())
    elements.append(t)
    elements.append(Spacer(1, 0.35 * inch))

    elements.append(Paragraph("Top 5 Autopartes más Vendidas", styles["Heading2"]))
    top_data = [["#", "SKU", "Nombre", "Cant. Vendida", "Ingresos"]]
    for i, part in enumerate(data["top_parts"], 1):
        top_data.append([
            str(i), part["sku"], part["name"],
            str(part["total_quantity_sold"]),
            f"${float(part['total_revenue']):,.2f}",
        ])
    t2 = Table(top_data, colWidths=[0.4*inch, 1.2*inch, 2.8*inch, 1.2*inch, 1.2*inch])
    t2.setStyle(_pdf_header_style())
    elements.append(t2)
    elements.append(Spacer(1, 0.35 * inch))

    elements.append(Paragraph("Pedidos por Estado", styles["Heading2"]))
    status_data = [["Estado", "Cantidad"]]
    for status, count in data["orders_by_status"].items():
        status_data.append([status.capitalize(), str(count)])
    t3 = Table(status_data, colWidths=[3 * inch, 3 * inch])
    t3.setStyle(_pdf_header_style())
    elements.append(t3)

    doc.build(elements)
    return buffer.getvalue()


# ── PDF: Recibo de un pedido individual ──────────────────────

def generate_order_pdf(order) -> bytes:
    """PDF/recibo de un pedido específico para descarga del cliente."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    elements = []

    short_id = str(order.id)[:8].upper()
    elements.append(Paragraph(f"Pedido #{short_id}", styles["Title"]))
    elements.append(Paragraph(f"Empresa: Autopartes MACUIN", styles["Normal"]))
    fecha = order.created_at.strftime("%Y-%m-%d %H:%M") if order.created_at else "N/A"
    elements.append(Paragraph(f"Fecha: {fecha}", styles["Normal"]))
    elements.append(Paragraph(f"Estado: {order.status.value.upper()}", styles["Normal"]))
    elements.append(Spacer(1, 0.35 * inch))

    elements.append(Paragraph("Detalle de Productos", styles["Heading2"]))
    items_data = [["SKU", "Producto", "Cant.", "Precio Unit.", "Subtotal"]]
    for item in order.items:
        sku  = item.part.sku  if item.part else "N/A"
        name = item.part.name if item.part else "N/A"
        items_data.append([
            sku, name,
            str(item.quantity),
            f"${float(item.unit_price):,.2f}",
            f"${float(item.subtotal):,.2f}",
        ])

    t = Table(items_data, colWidths=[1.0*inch, 2.5*inch, 0.7*inch, 1.4*inch, 1.2*inch])
    t.setStyle(_pdf_header_style())
    elements.append(t)
    elements.append(Spacer(1, 0.2 * inch))

    # Fila de total
    total_row = [["", "TOTAL", f"${float(order.total_amount):,.2f}"]]
    t2 = Table(total_row, colWidths=[4.2*inch, 1.5*inch, 1.1*inch])
    t2.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("BACKGROUND", (1, 0), (-1, -1), colors.HexColor("#27ae60")),
        ("TEXTCOLOR", (1, 0), (-1, -1), colors.white),
    ]))
    elements.append(t2)

    doc.build(elements)
    return buffer.getvalue()


# ── xlsx: Resumen de ventas ───────────────────────────────────

def generate_summary_xlsx(data: dict) -> bytes:
    """Excel con 3 hojas: Resumen, Top Autopartes, Por Estado."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    DARK = "2C3E50"
    BLUE = "2980B9"
    GREEN = "27AE60"

    def header_style(ws, row: int, fill_hex: str):
        fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
        font = Font(color="FFFFFF", bold=True)
        for cell in ws[row]:
            cell.fill = fill
            cell.font = font
            cell.alignment = Alignment(horizontal="center")

    wb = openpyxl.Workbook()

    # Hoja 1 — Resumen
    ws1 = wb.active
    ws1.title = "Resumen"
    ws1.append(["Reporte de Ventas — Autopartes API"])
    ws1.append([f"Generado: {_now_str()}"])
    ws1.append([])
    ws1.append(["Métrica", "Valor"])
    ws1.append(["Total de Ventas (excl. cancelados)", float(data["total_sales"])])
    ws1.append(["Total de Pedidos (excl. cancelados)", data["total_orders"]])
    header_style(ws1, 4, DARK)
    ws1.column_dimensions["A"].width = 38
    ws1.column_dimensions["B"].width = 18

    # Hoja 2 — Top Autopartes
    ws2 = wb.create_sheet("Top Autopartes")
    ws2.append(["#", "SKU", "Nombre", "Cantidad Vendida", "Ingresos ($)"])
    for i, part in enumerate(data["top_parts"], 1):
        ws2.append([i, part["sku"], part["name"],
                    part["total_quantity_sold"], float(part["total_revenue"])])
    header_style(ws2, 1, BLUE)
    for col, w in zip("ABCDE", [5, 14, 30, 18, 16]):
        ws2.column_dimensions[col].width = w

    # Hoja 3 — Por Estado
    ws3 = wb.create_sheet("Por Estado")
    ws3.append(["Estado", "Cantidad"])
    for status, count in data["orders_by_status"].items():
        ws3.append([status.capitalize(), count])
    header_style(ws3, 1, GREEN)
    ws3.column_dimensions["A"].width = 20
    ws3.column_dimensions["B"].width = 14

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ── docx: Resumen de ventas ───────────────────────────────────

def generate_summary_docx(data: dict) -> bytes:
    """Word con resumen, top partes y pedidos por estado."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    def set_cell_bg(cell, hex_color: str):
        """Colorea el fondo de una celda de tabla."""
        from docx.oxml import parse_xml
        shading = parse_xml(
            f'<w:shd {cell._element.nsmap["w"]} w:val="clear" '
            f'w:color="auto" w:fill="{hex_color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    doc = Document()
    doc.add_heading("Reporte de Ventas — Autopartes API", 0)
    doc.add_paragraph(f"Generado: {_now_str()}")

    # Resumen general
    doc.add_heading("Resumen General", level=1)
    t1 = doc.add_table(rows=3, cols=2)
    t1.style = "Table Grid"
    t1.cell(0, 0).text = "Métrica"
    t1.cell(0, 1).text = "Valor"
    t1.cell(1, 0).text = "Total de Ventas (excl. cancelados)"
    t1.cell(1, 1).text = f"${float(data['total_sales']):,.2f}"
    t1.cell(2, 0).text = "Total de Pedidos (excl. cancelados)"
    t1.cell(2, 1).text = str(data["total_orders"])

    # Top 5 autopartes
    doc.add_heading("Top 5 Autopartes más Vendidas", level=1)
    if data["top_parts"]:
        t2 = doc.add_table(rows=len(data["top_parts"]) + 1, cols=4)
        t2.style = "Table Grid"
        for j, h in enumerate(["SKU", "Nombre", "Cant. Vendida", "Ingresos"]):
            t2.cell(0, j).text = h
        for i, part in enumerate(data["top_parts"], 1):
            t2.cell(i, 0).text = part["sku"]
            t2.cell(i, 1).text = part["name"]
            t2.cell(i, 2).text = str(part["total_quantity_sold"])
            t2.cell(i, 3).text = f"${float(part['total_revenue']):,.2f}"
    else:
        doc.add_paragraph("Sin datos de ventas aún.")

    # Pedidos por estado
    doc.add_heading("Pedidos por Estado", level=1)
    statuses = list(data["orders_by_status"].items())
    t3 = doc.add_table(rows=len(statuses) + 1, cols=2)
    t3.style = "Table Grid"
    t3.cell(0, 0).text = "Estado"
    t3.cell(0, 1).text = "Cantidad"
    for i, (status, count) in enumerate(statuses, 1):
        t3.cell(i, 0).text = status.capitalize()
        t3.cell(i, 1).text = str(count)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── PDF: Reporte de pedidos ───────────────────────────────────

def generate_orders_pdf(data: dict) -> bytes:
    """PDF con tabla de pedidos recientes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Reporte de Pedidos — Autopartes API", styles["Title"]))
    elements.append(Paragraph(f"Generado: {_now_str()}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    rows = [["ID (corto)", "Estado", "Items", "Total ($)", "Fecha"]]
    for o in data["orders"]:
        rows.append([
            o["id"][:8].upper(),
            o["status"].capitalize(),
            str(o["items_count"]),
            f"${float(o['total_amount']):,.2f}",
            (o["created_at"] or "")[:10],
        ])

    t = Table(rows, colWidths=[1.1*inch, 1.2*inch, 0.8*inch, 1.3*inch, 1.3*inch])
    t.setStyle(_pdf_header_style())
    elements.append(t)
    elements.append(Spacer(1, 0.2 * inch))
    elements.append(Paragraph(f"Total de registros: {data['total']}", styles["Normal"]))

    doc.build(elements)
    return buffer.getvalue()


# ── PDF: Reporte de clientes ──────────────────────────────────

def generate_clients_pdf(data: dict) -> bytes:
    """PDF con ranking de clientes por volumen de compras."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Reporte de Clientes — Autopartes API", styles["Title"]))
    elements.append(Paragraph(f"Generado: {_now_str()}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    rows = [["#", "Nombre", "Email", "Pedidos", "Total Gastado ($)"]]
    for i, c in enumerate(data["clients"], 1):
        rows.append([
            str(i),
            c["full_name"],
            c["email"],
            str(c["total_orders"]),
            f"${float(c['total_spent']):,.2f}",
        ])

    t = Table(rows, colWidths=[0.4*inch, 2.0*inch, 2.2*inch, 0.9*inch, 1.3*inch])
    t.setStyle(_pdf_header_style())
    elements.append(t)
    elements.append(Spacer(1, 0.2 * inch))
    elements.append(Paragraph(f"Total de clientes: {data['total_clients']}", styles["Normal"]))

    doc.build(elements)
    return buffer.getvalue()


# ── PDF: Reporte de inventario ────────────────────────────────

def generate_inventory_pdf(data: dict) -> bytes:
    """PDF con stock actual de todas las autopartes activas."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Reporte de Inventario — Autopartes API", styles["Title"]))
    elements.append(Paragraph(f"Generado: {_now_str()}", styles["Normal"]))
    elements.append(Paragraph(
        f"Total partes: {data['total_parts']}  |  Stock bajo (≤5): {data['low_stock_count']}",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 0.3 * inch))

    rows = [["SKU", "Nombre", "Categoría", "Stock", "Precio ($)", "Alerta"]]
    for p in data["parts"]:
        rows.append([
            p["sku"],
            p["name"],
            p["category"] or "—",
            str(p["stock_quantity"]),
            f"${float(p['price']):,.2f}",
            "⚠ BAJO" if p["low_stock"] else "OK",
        ])

    t = Table(rows, colWidths=[1.0*inch, 2.2*inch, 1.2*inch, 0.7*inch, 1.0*inch, 0.8*inch])
    style = _pdf_header_style()
    # Colorea en rojo las filas con stock bajo
    for i, p in enumerate(data["parts"], 1):
        if p["low_stock"]:
            style.add("BACKGROUND", (0, i), (-1, i), colors.HexColor("#fadbd8"))
    t.setStyle(style)
    elements.append(t)

    doc.build(elements)
    return buffer.getvalue()
